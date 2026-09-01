#!/usr/bin/env python3
"""설계팀용 HOLD THE FLOW 제작 인계 DOCX를 생성한다."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs/20260901_HOLD_THE_FLOW_설계팀_제작인계_P0.docx"
ASSET_DIR = REPO_ROOT / "docs/assets/design_handoff"

FONT = "Noto Sans CJK KR"
INK = "111827"
MUTED = "667085"
BLUE = "0B5FFF"
BLUE_DARK = "173E8F"
BLUE_PALE = "EAF1FF"
GREEN = "087A55"
GREEN_PALE = "E8F7F1"
AMBER = "B54708"
AMBER_PALE = "FFF4E8"
RED = "B42318"
RED_PALE = "FDECEC"
LINE = "D9E2F0"
SLATE_PALE = "F6F8FB"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, size: float | None = None, *, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in CELL_MARGINS_DXA.items():
        tag = qn(f"w:{side}")
        node = margins.find(tag)
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def widths_from_weights(weights: list[float]) -> list[int]:
    total = sum(weights)
    widths = [round(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def apply_table_geometry(table, weights: list[float]) -> None:
    widths = widths_from_weights(weights)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Twips(widths[index])
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_paragraph_keep(paragraph, *, next_paragraph: bool = False, together: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    if together:
        node = OxmlElement("w:keepLines")
        p_pr.append(node)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("P0  ·  ")
    set_run_font(run, 8, bold=True, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text_node, end):
        run._r.append(node)


def apply_header_footer(section) -> None:
    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.clear()
    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        add_page_number(paragraph)


def add_picture(document: Document, filename: str, width: float, alt: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    shape = run.add_picture(str(ASSET_DIR / filename), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    set_paragraph_keep(paragraph, together=True)


def add_label(document: Document, text: str, color: str = BLUE) -> None:
    is_section_label = text[:2].isdigit()
    if is_section_label:
        spacer = document.add_paragraph()
        if getattr(document, "_hold_flow_page_break_pending", False):
            spacer.paragraph_format.page_break_before = True
            document._hold_flow_page_break_pending = False
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(4)
        run = spacer.add_run("\u00a0")
        set_run_font(run, 4, color=WHITE)
        set_paragraph_keep(spacer, next_paragraph=True, together=True)
    paragraph = document.add_paragraph(style="Eyebrow")
    if getattr(document, "_hold_flow_page_break_pending", False):
        paragraph.paragraph_format.page_break_before = True
        document._hold_flow_page_break_pending = False
    if is_section_label:
        paragraph.paragraph_format.space_before = Pt(18)
    paragraph.add_run(text.upper())
    paragraph.runs[0].font.color.rgb = RGBColor.from_string(color)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(paragraph, next_paragraph=True, together=True)


def add_body(document: Document, text: str, *, style: str = "Normal") -> None:
    paragraph = document.add_paragraph(text, style=style)
    set_paragraph_keep(paragraph, together=True)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="List Bullet")
        set_paragraph_keep(paragraph, together=True)


def add_checklist(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(item, style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(3)
        set_paragraph_keep(paragraph, together=True)


def add_numbered(document: Document, items: list[str]) -> None:
    numbering = document.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    next_num_id = max(int(item.get(qn("w:numId"))) for item in nums) + 1
    style_num_id = int(document.styles["List Number"]._element.pPr.numPr.numId.val)
    source = next(item for item in nums if int(item.get(qn("w:numId"))) == style_num_id)
    abstract_id = int(source.find(qn("w:abstractNumId")).get(qn("w:val")))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        paragraph = document.add_paragraph(item, style="List Number")
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = next_num_id
        set_paragraph_keep(paragraph, together=True)


def add_numbered_render_safe(document: Document, items: list[str]) -> None:
    """Render-safe numbered steps for LibreOffice page-break edge cases."""
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run(f"{index}.\u2003{item}")
        set_paragraph_keep(paragraph, together=True)


def add_callout(document: Document, title: str, body: str, *, tone: str = "amber") -> None:
    colors = {
        "amber": (AMBER, AMBER_PALE),
        "red": (RED, RED_PALE),
        "blue": (BLUE_DARK, BLUE_PALE),
        "green": (GREEN, GREEN_PALE),
    }
    ink, fill = colors[tone]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), fill)
        node.set(qn("w:space"), "8")
        borders.append(node)
    p_pr.append(borders)
    run = paragraph.add_run(title)
    set_run_font(run, 10, bold=True, color=ink)
    run.add_break()
    run = paragraph.add_run(body)
    set_run_font(run, 10, color=INK)
    set_paragraph_keep(paragraph, together=True)


def add_table(document: Document, headers: list[str], rows: list[list[str]], weights: list[float], *, font_size: float = 8.7) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_table_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        set_run_font(run, font_size, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, font_size, color=INK)
    apply_table_geometry(table, weights)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(document: Document) -> None:
    document._hold_flow_page_break_pending = True


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE_DARK, 18, 10),
        ("Heading 2", 13, INK, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.line_spacing = 1.05

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    subtitle.font.size = Pt(13)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    eyebrow = styles.add_style("Eyebrow", 1)
    eyebrow.font.name = FONT
    eyebrow._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    eyebrow.font.size = Pt(9)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = RGBColor.from_string(BLUE)
    eyebrow.paragraph_format.space_before = Pt(0)
    eyebrow.paragraph_format.space_after = Pt(8)
    eyebrow.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.space_before = Pt(8)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    apply_header_footer(section)


def build() -> None:
    document = Document()
    document.settings.odd_and_even_pages_header_footer = True
    configure_styles(document)
    configure_page(document)
    document.core_properties.title = "HOLD THE FLOW 설계팀 제작 인계 P0"
    document.core_properties.subject = "250 mm 이동형 양팔 물 따르기 로봇의 CAD·출력·URDF 제작 인계"
    document.core_properties.author = "HOLD THE FLOW Team"
    document.core_properties.keywords = "SO-101, URDF, K1 Max, bimanual robot, design handoff"

    # Cover
    document.add_paragraph().paragraph_format.space_after = Pt(12)
    add_label(document, "DESIGN HANDOFF · 2026.09.01")
    title = document.add_paragraph("HOLD THE FLOW", style="Title")
    title.paragraph_format.space_before = Pt(8)
    subtitle = document.add_paragraph("250 mm 이동형 양팔 물 따르기 로봇 · 제작 인계 P0", style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(16)
    add_picture(document, "01_p0_model_status.png", 6.5, "URDF P0의 운반 후보와 붓기 후보 자세 비교. 두 자세 모두 현재 설계 게이트를 통과하지 못했다.")
    add_callout(document, "현재 판정", "디지털 형상과 고정 TF는 생성 완료. 실물 치수, TCP/FK, 작업 자세 여유, 최종 관성, 슬라이서 프로파일은 아직 설계 동결 전이다.", tone="amber")
    add_page_break(document)

    # 1. Decision page
    add_label(document, "01 · RELEASE DECISION")
    add_heading(document, "이번 인계에서 확정한 것과 막아 둔 것", 1)
    add_body(document, "P0는 제작을 시작할 수 있는 기준 모델이지만, 곧바로 전체 부품을 출력하거나 관절값을 실기체에 전송하는 단계는 아니다. 아래 세 판정을 분리해 사용한다.")
    add_table(
        document,
        ["영역", "판정", "근거"],
        [
            ["CAD 형상", "조건부 통과", "차체 STL 14종 B-Rep 유효, K1 Max 수납, 45° 무서포트"],
            ["URDF 구조", "통과", "35 links / 34 joints, Xacro·check_urdf·메시·관성·축·limit 검사"],
            ["작업 자세", "차단", "TCP 없음, 붓기 hard limit 위반, 운반 링크·마스트 여유 미달"],
        ],
        [1.2, 1.0, 3.8],
    )
    add_heading(document, "고정할 수 있는 배치", 2)
    add_bullets(
        document,
        [
            "차체 외곽 250 × 250 mm, 3층 판 구조, JD-AMR형 2륜 + 후방 볼 캐스터.",
            "차체 중심 바닥 기준 팔 베이스 (25, ±70, 119) mm. 같은 점은 base_footprint 기준 (-65, ±70, 119) mm.",
            "Astra S 광학 중심은 차체 중심 기준 (-80, 0, 800) mm, LDS-03은 (90, 0, 165) mm.",
            "운반 중 팔 자세 고정, 붓기 중 이동 베이스 정지. 식탁은 요구하지 않고 낮은 작업 패드를 사용.",
        ],
    )
    add_callout(document, "STOP", "현재 YAML의 transport_joint_degrees와 pour_joint_degrees는 실기체로 보내지 않는다. TCP/FK와 여유 검증을 통과한 새 자세만 제어 입력으로 승인한다.", tone="red")
    add_page_break(document)

    # 2. URDF
    add_label(document, "02 · URDF REVALIDATION")
    add_heading(document, "URDF는 구조 PASS, 물 따르기 task는 FAIL", 1)
    add_body(document, "검증은 구조·수치 품질·작업 자세·충돌 여유 네 층으로 나눴다. 구조 PASS를 작업 가능 PASS로 해석하지 않는다.")
    add_table(
        document,
        ["검사층", "결과", "핵심 수치"],
        [
            ["Xacro/트리", "PASS", "35 links, 34 joints, Xacro=커밋 URDF, check_urdf 통과"],
            ["메시/수치", "PASS", "51 refs / 24 files, 33 양의 질량·SPD 관성, 16 정규축, 14 hard limit"],
            ["TCP/FK", "FAIL", "left/right tool0와 bottle/cup TCP 4개 누락"],
            ["붓기 관절", "FAIL", "양쪽 link2_to_link3가 0~3.316126 rad 범위 밖"],
            ["운반 여유", "FAIL", "링크 18.78<25 mm, 마스트 반경 31.36<55 mm"],
        ],
        [1.3, 0.8, 3.9],
    )
    add_heading(document, "붓기 후보의 hard limit 위반", 2)
    add_table(
        document,
        ["관절", "후보값", "URDF 범위", "판정"],
        [
            ["left_link2_to_link3", "-0.396190 rad", "0~3.316126 rad", "FAIL"],
            ["right_link2_to_link3", "-0.122173 rad", "0~3.316126 rad", "FAIL"],
        ],
        [2.0, 1.2, 1.5, 0.8],
    )
    add_heading(document, "근본 원인", 2)
    add_body(document, "작업 자세는 JD-AMR에서 검증한 SO-101 관절 체인과 TCP 기준에서 나왔지만, 현재 URDF에는 Robonine 전체 팔 Xacro가 들어갔다. 두 모델의 joint origin·axis·zero·limit가 달라 좌표와 관절값이 같은 기준이 아니다.")
    add_callout(document, "권고 아키텍처", "JD-AMR/SO-101 팔 체인은 유지하고 Robonine의 link5 이후 평행 죠 기구만 엔드이펙터로 부착한다. 이후 좌우 tool0/TCP, 서보 zero·sign·offset, 동일 URDF 기반 IK/FK 회귀시험을 만든다.", tone="blue")
    add_page_break(document)

    # 3. layout
    add_label(document, "03 · MODEL & COORDINATES")
    add_heading(document, "좌표계부터 맞추고 체결공은 실측 뒤 확정", 1)
    add_picture(document, "03_layout_reference.png", 6.5, "250 mm 차체의 상면과 측면 URDF 배치. 차체 중심 기준 팔, LiDAR, 카메라, 바퀴, 캐스터 좌표를 표시했다.")
    add_body(document, "배치도는 차체 중심 바닥을 원점으로 사용한다. ROS TF의 base_footprint는 구동 바퀴축 바닥 투영점이므로 차체 중심보다 X 방향으로 90 mm 앞이다. 설계 도면과 URDF 수치를 비교할 때 X 변환을 반드시 적용한다.")
    add_table(
        document,
        ["대상", "차체 중심 바닥 기준 mm", "base_footprint 기준 mm"],
        [
            ["왼팔 / 오른팔", "(25, ±70, 119)", "(-65, ±70, 119)"],
            ["Astra optical", "(-80, 0, 800)", "(-170, 0, 800)"],
            ["LDS-03", "(90, 0, 165)", "(0, 0, 165)"],
            ["바퀴축", "X=+90, Y=±135", "X=0, Y=±135"],
            ["후방 캐스터", "X=-105", "X=-195"],
        ],
        [1.3, 2.2, 2.2],
    )
    add_page_break(document)

    # 4. Part catalog
    add_label(document, "04 · PRINT PACKAGE")
    add_heading(document, "K1 Max 출력 대상 14종", 1)
    add_picture(document, "02_print_part_catalog.png", 6.5, "K1 Max 원본 방향 출력 대상인 차체 CAD 14종의 형상, 크기, 수량 목록.")
    add_callout(document, "범위 주의", "이 목록은 차체 CAD다. Robonine 평행그리퍼의 Main frame·Clamp·Holder는 별도 검증 트리이며 차체 14종에 포함되지 않는다.", tone="amber")

    add_label(document, "05 · PRINT BOM")
    add_heading(document, "출력 순서와 동결 조건", 1)
    add_table(
        document,
        ["부품 / 수량", "크기 mm", "첫 출력", "동결 전 확인"],
        [
            ["하판 / 1", "250×250×48", "M1 뒤", "휨·대각선·휠 캐리어 동시 천공"],
            ["중판 / 1", "250×250×6", "하판 뒤", "배터리·SBC 외피와 케이블홀"],
            ["상판 / 1", "250×250×8", "건식 배치 뒤", "팔·마스트 보강판 관통 체결"],
            ["팔 어댑터 / 2", "130×90×6", "시험 1개", "SO-101 두 대 홀 패턴 실측"],
            ["마스트 보강판 / 1", "75×75×8", "시험 1개", "M6 관통과 너트 접근"],
            ["마스트 세그먼트 / 3", "45×35×267", "1개", "60 mm 겹침·직각도·비틀림"],
            ["마스트 이음관 / 2", "38.2×28.2×60", "1개", "0.20~0.35 mm 끼움 쿠폰"],
            ["Astra 거치대 / 1", "65×182×30", "실측 뒤", "M6·커넥터·피치 35~45°"],
            ["LiDAR 받침 / 1", "65×50×29", "실측 뒤", "하부 홀·스캔면 Z≈165"],
            ["휠 캐리어 / 2", "82×58×8", "시험 1개", "혼·625 베어링·축 스택"],
            ["캐스터 어댑터 / 1", "70×52×8", "실측 뒤", "플랜지 홀·접지 높이"],
            ["시임 / 각 1", "70×52×1/2/3", "필요 시", "세 접지점 높이 차 ≤0.5"],
        ],
        [1.25, 1.05, 1.05, 2.65],
        font_size=8.2,
    )
    add_callout(document, "G-code 승인 조건", "M0 실물 측정표, M1 장공·끼움 쿠폰과 판 공차·휨 결과, 실제 K1 Max/노즐/PETG 3MF와 슬라이서 preview가 모두 있어야 한다. 현재 저장소에는 양산용 G-code가 없다.", tone="red")
    add_page_break(document)

    # 6. Printer process
    add_label(document, "06 · K1 MAX WORKFLOW")
    add_heading(document, "한 번에 전체를 뽑지 않는 출력 절차", 1)
    add_numbered(
        document,
        [
            "프린터 실사용 베드, 프라임 라인, 제외영역을 확인한다. 250 mm 판과 최대 8 mm 브림이 겹치면 프라임 위치 또는 브림을 조정한다.",
            "실물 SO-101, Astra S, LDS-03, 캐스터, C018 혼·625 베어링·축을 캘리퍼로 재고 사진과 측정표를 남긴다.",
            "M1 장공·끼움·너트 포켓 쿠폰을 출력해 XY 홀 보정과 0.20/0.25/0.30/0.35 mm 끼움 여유를 선택한다.",
            "팔 어댑터, 마스트 세그먼트·이음관, 휠 캐리어를 각각 1개만 시험 출력한다.",
            "M1 통과 뒤 하판 1장을 출력한다. 250±0.5 mm, 대각선 차 ≤1 mm, 휨 ≤1 mm를 만족하지 않으면 중·상판을 출력하지 않는다.",
            "하판 건식 조립 뒤 중판·상판을 순서대로 출력하고, 마지막에 센서 거치대와 시임을 확정한다.",
        ],
    )
    add_heading(document, "시작 프로파일", 2)
    add_table(
        document,
        ["항목", "P0 시작값", "판정 방법"],
        [
            ["재료", "PETG", "실제 롤 제조사·색상·건조 상태 기록"],
            ["노즐 / 레이어", "0.4 mm / 0.20 mm", "쿠폰과 판 표면·홀 치수 비교"],
            ["벽 / 상하부", "6 walls / 6 top / 6 bottom", "팔·마스트 하중 경로와 홀 주변 밀폐"],
            ["인필", "40~50% gyroid", "슬라이서 질량·판 휨 결과로 조정"],
            ["서포트", "OFF", "45° 분석 결과. Robonine 그리퍼는 별도"],
            ["브림", "큰 판 6~8 mm", "베드 제외영역·프라임 라인 확인"],
        ],
        [1.3, 1.8, 2.9],
    )
    add_page_break(document)

    # 7. Assembly
    add_label(document, "07 · ASSEMBLY")
    add_heading(document, "조립은 바닥에서 위로, 케이블은 움직임보다 먼저", 1)
    add_numbered_render_safe(
        document,
        [
            "후방 캐스터와 좌우 휠 모듈을 하판에 임시 체결하고 세 접지점의 흔들림과 바닥 간극을 측정한다.",
            "배터리·전원 분배·SBC·제어보드를 하판/중판 사이에 건식 배치하고 탈착 방향과 서비스 루프를 확인한다.",
            "중판을 체결한 뒤 LiDAR 받침을 설치해 360° 스캔면의 차폐를 확인한다.",
            "상판, 좌우 팔 어댑터, 하부 보강판을 관통 체결한다. 인쇄물 나사산에 반복 체결하지 않는다.",
            "중앙 후방 마스트를 3세그먼트로 조립하고 각 이음부 60 mm 겹침, 직각도, 비틀림을 기록한다.",
            "Astra S, 평행그리퍼, 케이블 스트레인 릴리프를 설치하고 zero/운반 후보 수동 스윕으로 간섭을 확인한다.",
        ],
    )
    add_heading(document, "체결 원칙", 2)
    add_bullets(
        document,
        [
            "차체 판·팔·마스트는 관통 볼트 + 와셔 + 나일록 너트 또는 열압입 인서트 사용.",
            "휠 허브는 C018 혼–JD-AMR 휠–625 베어링–축–와셔–스페이서의 축방향 스택 도면을 먼저 확정.",
            "볼트 길이 산정식: 체결물 두께 + 와셔 + 너트 유효나사 2~3산. 바닥 돌출과 회전체 간섭 금지.",
            "Astra USB와 팔 버스 케이블은 최소 굽힘반경, 커넥터 인출 공간, 정비용 서비스 루프를 확보.",
        ],
    )
    add_callout(document, "전원 안전", "모든 팔 서보 라벨이 STS3215-C018인지 확인하기 전에는 12 V 팔 버스를 연결하지 않는다. C001이 한 개라도 섞이면 해당 버스 전압을 다시 설계한다.", tone="red")
    add_page_break(document)

    # 8. Mass and operation
    add_label(document, "08 · MASS, MOMENT & MOTION")
    add_heading(document, "무게 계산은 가능성을 보여 주지만 실물 합격을 대신하지 않는다", 1)
    add_table(
        document,
        ["검토값", "P0 계산", "사용 조건"],
        [
            ["총질량", "약 5.26~5.34 kg", "실물 부품·배선·체결재 포함 후 재측정"],
            ["주행 토크/휠", "0.234~0.237 N·m (SF 2)", "적재 가정별 값. C018 정격 0.981 N·m과 실물 시험"],
            ["붓기 COM", "x=36.7, y=11.5, z=247.2 mm", "120 g 물체 + Robonine 기준 그리퍼 가정"],
            ["동적 전방 여유", "24.1 mm", "5° 경사 + 0.3 m/s² 감속 가정"],
            ["컵 유지 토크", "0.948 N·m", "120 g 물체. 손목 전류·온도 기록"],
            ["140 g 유지 토크", "1.003 N·m", "Robonine 원형 기준 정격 초과, 금지"],
        ],
        [1.6, 1.7, 2.7],
    )
    add_heading(document, "가동범위 검증 순서", 2)
    add_numbered(
        document,
        [
            "MoveIt/Isaac Sim에서 self-collision matrix를 만들되, zero joint pose를 홈 자세로 사용하지 않는다.",
            "새 tool0/TCP 기준으로 운반·접근·파지·붓기·배치 자세를 IK로 다시 푼다.",
            "각 trajectory를 연속 충돌검사하고 최소 링크 25 mm, 그리퍼-반대팔 40 mm, 마스트 반경 55 mm를 적용한다.",
            "실물은 무부하 10% 속도에서 시작해 120 g 물병과 빈 컵으로 확장한다. 누수 트레이와 E-stop을 준비한다.",
        ],
    )
    add_callout(document, "현재 금지", "가득 찬 500 mL 병, 300 g 물체, 움직이는 베이스에서 붓기, 임시 부호 반전으로 관절 한계를 우회하는 동작은 P0 범위 밖이다.", tone="red")
    add_page_break(document)

    # 9. Evidence matrix
    add_label(document, "09 · EVIDENCE GATES")
    add_heading(document, "설계 완료는 파일이 아니라 증거 묶음으로 판단", 1)
    add_table(
        document,
        ["Gate", "필수 증거", "통과 기준"],
        [
            ["M0 실측", "CSV+기준점 사진+서보 라벨", "모든 미확정 체결부 측정 완료"],
            ["M1 공차", "쿠폰 사진+선택값+3MF", "홀·인서트·끼움·로드 기준 선택"],
            ["M2 건식조립", "상/측/하 사진+대각선·뒤틀림", "정사각도·평탄도·3점 지지 통과"],
            ["M3 정적하중", "1 N 변위 영상+10분 유지", "균열·영구변형·이음 미끄럼 없음"],
            ["M4 전도", "3점 저울표+5° 경사 영상", "각 접지 5% 이상, 여유 20 mm 이상"],
            ["M5 전원", "전압·전류 로그+E-stop 영상", "모터 차단·SBC 로그 유지"],
            ["M6 단팔", "온도·전류·미끄럼·FSR 로그", "120 g 통과 뒤 다음 하중 진행"],
            ["M6D 주행", "5 m 직진+회전 10회+20분 로그", "55°C 미만, 허브 풀림 없음"],
            ["M7 양팔", "100회 trajectory+저속 영상", "충돌 0, 중앙영역 lease 동작"],
            ["M8 센서", "유효 depth+가림각+costmap", "task 시야와 주행 스캔 통과"],
            ["M9 URDF/Sim", "TF/FK 오차표+축·충돌 로그", "CAD·실물·시뮬 기준 일치"],
        ],
        [0.8, 2.7, 2.5],
        font_size=8.3,
    )
    add_heading(document, "파일 단일 원본", 2)
    add_table(
        document,
        ["대상", "저장소 경로"],
        [
            ["기계 명세", "design/mechanical/hold_flow_mechanical_v0_2.yaml"],
            ["CAD 생성", "design/cad/generate_hold_flow_cad.py"],
            ["출력 manifest", "design/cad/exports/manifest.json"],
            ["로봇 설명", "src/hold_flow_description/urdf/hold_flow.urdf.xacro"],
            ["실물 보정", "src/hold_flow_description/config/mechanical_calibration.yaml"],
            ["URDF 검증", "validate_description.py / audit_urdf_quality.py / audit_task_pose.py"],
            ["간섭 검증", "design/cad/audit_urdf_clearance.py"],
        ],
        [1.3, 4.7],
        font_size=8.4,
    )
    # 10. Checklist
    add_label(document, "10 · RELEASE CHECKLIST")
    add_heading(document, " 설계팀 최종 체크", 1)
    add_heading(document, "출력 전", 2)
    add_checklist(
        document,
        [
            "□ 실물 부품 라벨과 캘리퍼 측정표가 YAML에 반영됨",
            "□ M1 쿠폰으로 홀·장공·끼움·인서트 공차를 선택함",
            "□ K1 Max 실제 프로파일과 3MF preview를 저장함",
            "□ 250 mm 판과 브림이 베드·프라임 라인·제외영역 안에 있음",
        ],
    )
    add_heading(document, "URDF 동결 전", 2)
    add_checklist(
        document,
        [
            "□ JD-AMR 팔 체인 + Robonine 평행 죠 결합 방식을 확정함",
            "□ left/right tool0와 bottle/cup TCP 4개가 존재함",
            "□ 서보 raw↔URDF zero·sign·offset을 실물로 대조함",
            "□ 운반·붓기·배치 자세를 동일 URDF에서 다시 산출함",
            "□ 최소 링크 25 mm, 그리퍼 40 mm, 마스트 55 mm 여유를 전 구간 통과함",
        ],
    )
    add_heading(document, "전원·통합 전", 2)
    add_checklist(
        document,
        [
            "□ 모든 팔 서보의 C001/C018 변형을 확인함",
            "□ E-stop이 좌우 팔과 주행 전원을 끊고 SBC 로그는 유지함",
            "□ 120 g 물병과 빈 컵으로 10% 속도 반복시험을 통과함",
            "□ 누수 트레이, 미끄럼 감지, 실패 복구 상태를 준비함",
        ],
    )
    add_callout(document, "P0 종료 조건", "M0~M9와 M6D 증거가 모이고 task pose audit와 clearance audit가 모두 PASS가 되면 설계 동결 후보로 올린다. 그전까지는 형상과 치수 변경 이력을 남기는 P0 단계다.", tone="green")
    add_body(document, "상세 명세와 명령은 저장소의 docs/20260901_설계팀_제작인계패키지_P0.md를 기준으로 한다.", style="Caption")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
